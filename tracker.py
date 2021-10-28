'''
    BREAK DOWN OF HOW THE TRACKER WILL WORK
        1. READ DATA FROM ORIGINAL FILE
        2. OPTION TO UPDATE BY SELECTING FILE
        2. HAVE A UI JUST TO CHOOSE THE FILE
        3. SAVES THE FILE AND ITS FORMAT
'''
import pandas as pd
import time as t

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Pt
from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

tmr_dict = {}

week = []
two_weeks = []
later_weeks = []
        
def importing_logic():
    # reads excel file
    page = pd.read_excel('tmr page.xls')
    # converts file to a list
    tmr_list = list(page.values.tolist())
    # removing non tmr lines
    tmr_list.pop(-1)
    tmr_list.pop(0)

    for i in tmr_list:
        name = i[9]
        num = i[1]
        start_date = i[3]
        sd = start_date[:10]
        st = start_date[10:16]
        start = f'{sd} @{st}'
        end_date = i[5]
        ed = end_date[:10]
        et = end_date[10:16]
        end = f'{ed} @{et}'
        time = f'{start} -> {end}'
        sl = i[2]
        el = i[4]
        location = f'{sl} -> {el}'
        su = i[8]
        status = i[10]
        unit_status = f'{su}: {status}'

        tmr_dict[i[9]] = [name, num, time, location, unit_status]


def create_doc():
    document = Document()

    section = document.sections[0]
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'),'2')

    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(0.5)
    section.right_margin = Cm(0.5)

    p = document.add_paragraph(f'Updated: {t.ctime()}')

    for k, v in tmr_dict.items():
        content = f'{v[1]}: {v[0]}\n{v[2]}\n{v[3]} // {v[4]}\nRequirements:________________________________________________________________________________________'
        p = document.add_paragraph().add_run(content)
        p.font.size = Pt(6)

    document.save('results.docx')

importing_logic()
print(f'\nIMPORTED AT {t.ctime()}\n')
print(f'current time is {t.da}')

create_doc()